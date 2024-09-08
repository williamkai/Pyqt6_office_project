# order_sorting_window.py
import os
from PyQt6.QtCore import pyqtSignal, QTimer
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from collections import defaultdict
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt6.QtWidgets import (QWidget, QTableWidget, QTableWidgetItem, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QListWidget, QMessageBox)

class OrderSortingWindow(QWidget):
    closed = pyqtSignal()
    data_changed = pyqtSignal(dict)

    def __init__(self, parent=None, processed_data=None,folder_path=None):
        super().__init__(parent)  
        self.processed_data = processed_data
        self.folder_path=folder_path
        self.order_list = None
        self.order_table = None
        self.stats=None
        self.setWindowTitle("訂單處理結果")
        self.setMinimumSize(300, 400)
        self.resize(400, 700)
        self.initialize_ui()
        self.create_order_list()
        self.calculate_statistics()
        self.show()
        QTimer.singleShot(1000, self.merge_and_send_data)  

    def initialize_ui(self):
        self.main_layout = QVBoxLayout(self)
        self.display_area = QWidget()
        self.display_layout = QVBoxLayout(self.display_area)
        self.but_layout =QHBoxLayout()      
        
        self.create_order_list_widget()
        self.create_order_table_widget()
        self.create_save_button()
        
        self.main_layout.addWidget(self.display_area)

    def create_order_list_widget(self):
        self.order_list = QListWidget(self)
        self.order_list.setMinimumHeight(50)
        self.order_list.itemClicked.connect(self.display_order_table)
        self.display_layout.addWidget(self.order_list, 1)

    def create_order_table_widget(self):
        self.order_table = QTableWidget(self)
        self.display_layout.addWidget(self.order_table, 7)

    def create_save_button(self):
        self.save_button = QPushButton("Save to Word", self)
        self.save_button.clicked.connect(self.save_table_to_word)
        self.but_layout.addWidget(self.save_button)
        self.display_layout.addLayout(self.but_layout)

    def create_order_list(self):
        order_list_title = ["大榮", "通盈", "大榮商品總數", "通盈商品總數"]
        self.order_list.addItems(order_list_title)

    def calculate_statistics(self):
        self.stats = defaultdict(lambda: defaultdict(int))
        for category in ['大榮', '通盈']:
            category_data = self.processed_data.get(category, {})
            for order_id, order_details in category_data.items():
                for item in order_details['items']:
                    product_code = item['product_code']
                    quantity = int(item['quantity'])
                    self.stats[category][product_code] += quantity
        print(f"{self.stats}")
    
    """這邊暫時留著因為使用defaultdict，字典使用方式我還不清楚，如果後續要使用到總計數量
    可能會出問題，以防萬一，我沒辦法理解defaultdict，我就回歸原本字典方式使用
    
    """
    # def calculate_statistics(self):
    #     self.stats = {'大榮': {}, '通盈': {}}
    #     for category in ['大榮', '通盈']:
    #         category_data = self.processed_data.get(category, {})
    #         for order_id, order_details in category_data.items():
    #             for item in order_details['items']:
    #                 product_code = item['product_code']
    #                 quantity = int(item['quantity'])
    #                 if product_code in self.stats[category]:
    #                     self.stats[category][product_code] += quantity
    #                 else:
    #                     self.stats[category][product_code] = quantity

    def display_order_table(self):
        self.order_table.clear()
        selected_item = self.order_list.currentItem()
        if not selected_item:
            return
        selected_key = selected_item.text()
        if selected_key in ["大榮商品總數", "通盈商品總數"]:
            cleaned_key = selected_key.replace("商品總數", "")
            self.create_order_total_table(cleaned_key)
        else:
            data_to_display = self.processed_data.get(selected_key, {})
            self.create_order_table(data_to_display)

    def create_order_table(self, data):
        data = self.convert_keys_to_str(data)
        self.order_table.setRowCount(0)
        self.order_table.setColumnCount(4)
        self.order_table.setHorizontalHeaderLabels(['訂單來源', '商品代號', '數量', '單位'])
        
        row_count = 0
        for key, order in data.items():
            for i, item in enumerate(order['items']):
                self.order_table.insertRow(row_count)
                if i == 0:
                    self.order_table.setItem(row_count, 0, QTableWidgetItem(order['title']))
                self.order_table.setItem(row_count, 1, QTableWidgetItem(item['product_code']))
                self.order_table.setItem(row_count, 2, QTableWidgetItem(item['quantity']))
                self.order_table.setItem(row_count, 3, QTableWidgetItem('箱'))
                row_count += 1

        self.order_table.resizeColumnsToContents()
        self.order_table.resizeRowsToContents()

    def create_order_total_table(self,cleaned_key):
        data_to_display = self.stats.get(cleaned_key, {})
        self.order_table.setRowCount(0)
        self.order_table.setColumnCount(3)
        self.order_table.setHorizontalHeaderLabels(['商品代號', '數量', '單位'])
        for row_count, (code, total_quantity) in enumerate(data_to_display.items()):
            self.order_table.insertRow(row_count)
            self.order_table.setItem(row_count, 0, QTableWidgetItem(code))
            self.order_table.setItem(row_count, 1, QTableWidgetItem(str(total_quantity)))
            self.order_table.setItem(row_count, 2, QTableWidgetItem('箱'))
            row_count += 1
        self.order_table.resizeColumnsToContents()
        self.order_table.resizeRowsToContents()

    def convert_keys_to_str(self, data):
        if isinstance(data, dict):
            return {str(k): self.convert_keys_to_str(v) for k, v in data.items()}
        elif isinstance(data, list):
            return [self.convert_keys_to_str(i) for i in data]
        else:
            return data

    def save_table_to_word(self):
        selected_item = self.order_list.currentItem()
        if not selected_item:
            return

        try:
            selected_key = selected_item.text()
            if selected_key in ["大榮商品總數", "通盈商品總數"]:
                self.handle_total_statistics()
            else:
                self.handle_order_details(selected_item)
        except Exception as e:
            QMessageBox.critical(self, "錯誤", f"儲存至 Word 檔案時發生錯誤: {e}")
                
    def handle_order_details(self, selected_item):
        doc = Document()
        row_count = self.order_table.rowCount()
        for row in range(row_count):
            title = self.order_table.item(row, 0).text() if self.order_table.item(row, 0) else ''
            product_code = self.order_table.item(row, 1).text() if self.order_table.item(row, 1) else ''
            quantity = self.order_table.item(row, 2).text() if self.order_table.item(row, 2) else ''
            unit = self.order_table.item(row, 3).text() if self.order_table.item(row, 3) else ''

            if title:
                if selected_item.text() == '大榮':
                    paragraph = doc.add_paragraph(f"{title} 大榮-久富餘-菁弘")
                else:
                    paragraph = doc.add_paragraph(f"{title} 通盈-久富餘-菁弘")

                self.set_font(paragraph, size=14, font_name='標楷體')

            paragraph = doc.add_paragraph(f"{product_code}……{quantity}{unit}")
            self.set_font(paragraph, size=14, font_name='標楷體')

        today_date = datetime.now().strftime('%m%d')
        year = datetime.now().year - 1911
        today_date_tw = f"{year}{today_date}"
        file_name = f'{today_date_tw}{selected_item.text()}.docx'
        file_path = self.get_full_file_path(file_name)
        try:
            doc.save(file_path)
            print(f"File saved to {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "錯誤", f"儲存檔案時發生錯誤: {e}")

    def handle_total_statistics(self):
        doc = Document()
        # 獲取今天的日期並格式化
        today_date = datetime.now().strftime('%Y/%m/%d')
        
        # 在文件開頭添加日期並設置置中對齊
        date_paragraph = doc.add_paragraph(f"日期: {today_date}", style='Normal')
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for key, data in self.stats.items():
            heading = doc.add_heading(key, level=1)
            self.set_font(heading, size=26, font_name='標楷體', bold=True)

            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Shading Accent 5"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '商品代號'
            hdr_cells[1].text = '總箱數'
            self.set_font_cells(hdr_cells, size=22)

            for item_code, quantity in data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = item_code
                row_cells[1].text = f'{quantity}'
                self.set_font_cells(row_cells, size=18)

        today_date = datetime.now().strftime('%m%d')
        year = datetime.now().year - 1911
        today_date_tw = f"{year}{today_date}"
        file_name = f'{today_date_tw}統計數量.docx'
        file_path = self.get_full_file_path(file_name)
        try:
            doc.save(file_path)
            print(f"File saved to {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "錯誤", f"儲存檔案時發生錯誤: {e}")

    def set_font(self, paragraph, size, font_name, bold=False):
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.bold = bold

    def set_font_cells(self, cells, size):
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(size)
                run.font.name = '標楷體'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                run.font.color.rgb = RGBColor(0, 0, 0)

    def get_full_file_path(self, file_name):
        if not self.folder_path.endswith(os.path.sep):
            self.folder_path += os.path.sep
        return os.path.join(self.folder_path, file_name)

    def merge_and_send_data(self):
        combined_data = {
            '訂單資料': self.processed_data,
            '總計資料': self.stats
            }
        self.data_changed.emit(combined_data)

    def closeEvent(self, event):
        self.closed.emit()  # 當視窗關閉時發出訊號
        super().closeEvent(event)