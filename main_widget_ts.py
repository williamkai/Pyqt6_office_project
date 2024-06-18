import imaplib
import shutil
import email
import os

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
from email.header import decode_header, make_header
from PyQt6.QtWidgets import (
                            QWidget, 
                            QVBoxLayout, 
                            QHBoxLayout, 
                            QPushButton,
                            QLabel,
                            QDateEdit,
                            QTextEdit,
                            QLineEdit,
                            QCheckBox,QListWidget, QSpacerItem, QSizePolicy,QDialogButtonBox, QTextBrowser, QListWidgetItem,QDialog,QTimeEdit)

from PyQt6.QtGui import QPixmap
from PyQt6.QtCore import QThread, pyqtSignal,Qt, QTime
from PyQt6 import QtCore 

class MainWidget(QWidget):
    def __init__(self, parent, email, password):
        super().__init__(parent)
        self.email = email
        self.password = password
        self.mail_ids=''
        self.emails = []
        self.mail_ids = []
        self.order_search_datetime=""

        self.main_layout = QVBoxLayout(self)
        # 创建水平布局
        self.button_layout = QHBoxLayout()
        self.horizontal_layout = QHBoxLayout()
        self.horizontal_layout_1= QHBoxLayout()
        self.from_input_layout=QHBoxLayout()

        self.order_search_button = QPushButton("訂單搜索功能", self)
        self.temp_button1 = QPushButton("訂單整理功能", self)
        self.temp_button2 = QPushButton("轉寄信件功能", self)

        self.order_search_button.setFixedWidth(100)
        self.temp_button1.setFixedWidth(100)
        self.temp_button2.setFixedWidth(100)
        self.button_layout.addWidget(self.order_search_button)
        self.button_layout.addWidget(self.temp_button1)
        self.button_layout.addWidget(self.temp_button2)

        # 添加日期选择器
        self.date_edit = QDateEdit(self)
        self.date_edit.setCalendarPopup(True)
        self.date_edit.setDisplayFormat("yyyy-MM-dd")
        self.date_edit.setDate(datetime.today())
        self.date_edit.setFixedWidth(150)
        self.horizontal_layout_1.addWidget( self.date_edit)
        #添加時間選擇器
        self.timeedit = QTimeEdit(self)
        self.timeedit.setDisplayFormat('hh:mm:ss')
        self.time = QTime(12, 0, 0)  # 参数依次是小时、分钟、秒
        self.timeedit.setTime(self.time)
        self.timeedit.setFixedWidth(150)
        self.horizontal_layout_1.addWidget( self.timeedit)
        spacer_item = QSpacerItem(0, 0, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        self.horizontal_layout_1.addItem(spacer_item)
        self.main_layout.addLayout(self.horizontal_layout_1)

        # 添加一个 QLineEdit 用于输入 from 地址
        self.from_input = QLineEdit(self)
        self.from_input.setPlaceholderText("输入寄件人郵箱地址")
        self.from_input.setFixedWidth(200)
        self.from_input_layout.addWidget( self.from_input)
        #輸入收件人的郵件地址
        self.from_input_1 = QLineEdit(self)
        self.from_input_1.setPlaceholderText("输入收件人郵箱地址")
        self.from_input_1.setFixedWidth(200)
        self.from_input_layout.addWidget( self.from_input_1)
        spacer_item = QSpacerItem(0, 0, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        self.from_input_layout.addItem(spacer_item)
        self.main_layout.addLayout(self.from_input_layout)

        # 添加一个复选框来选择是否使用 from 地址作为筛选条件
        self.use_from_checkbox = QCheckBox("使用寄件人做篩選條件", self)
        self.main_layout.addWidget(self.use_from_checkbox)

         # 添加关键词输入框及标签
        self.keyword_label = QLabel("内容關鍵字篩選:", self)
        self.keyword_input = QLineEdit(self)
        self.keyword_input.setFixedWidth(200)
        self.keyword_input.setPlaceholderText("输入郵件關鍵詞")
        self.keyword_input.setText("久富餘-菁弘")  # 默认关键词
        self.horizontal_layout.addWidget( self.keyword_label)
        self.horizontal_layout.addWidget( self.keyword_input)
        # 添加水平空間，以便控件靠左排列
        spacer_item = QSpacerItem(0, 0, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        self.horizontal_layout.addItem(spacer_item)

        # 将水平布局管理器添加到主垂直布局管理器中
        self.main_layout.addLayout(self.horizontal_layout)
       
        #信件標題清單
        self.email_list = QListWidget(self)
        self.email_list.setMinimumHeight(50)
        self.main_layout.addWidget(self.email_list,1)

        # 添加一个 QTextEdit 来显示邮件内容
        self.email_display = QTextEdit(self)
        self.email_display.setReadOnly(True)
        self.email_display.setMinimumHeight(200)
        self.main_layout.addWidget(self.email_display,4)

        self.main_layout.addLayout(self.button_layout)

        self.setLayout(self.main_layout)

        self.order_search_button.clicked.connect(self.order_search)
        self.email_list.itemClicked.connect(self.display_email)
        self.temp_button1.clicked.connect(self.order2word)
        self.temp_button2.clicked.connect(self.send_email)

    def order_search(self):
        print("測試測試-訂單功能")
        print(f"使用的邮箱: {self.email}")

        # 取得使用者選擇的日期
        selected_date = self.date_edit.date().toString("dd-MMM-yyyy")
        selected_time = self.timeedit.time().toString("hh:mm")

        print(f"選擇的日期:{selected_date}\n選擇的時間:{selected_time}")
        # self.order_search_datetime_1=str(int(self.date_edit.date().toString("yyyy"))-1911)
        self.order_search_datetime=str(int(self.date_edit.date().toString("yyyy"))-1911)+self.date_edit.date().toString("MMdd")

        # 获取 from 地址输入框的值和复选框的状态
        from_address = self.from_input.text()
        use_from_condition = self.use_from_checkbox.isChecked()

        # 构建搜索条件
        search_criteria =f'(ON {selected_date})'#f'(SINCE {selected_date} 00:00:00 BEFORE {selected_date} 12:00:00)' 
        if use_from_condition and from_address:
            search_criteria += f' FROM "{from_address}"'

        print(f"搜索条件: {search_criteria}")

        # 获取关键词
        keyword = self.keyword_input.text()
        print("Keyword:", keyword)


        self.worker = EmailWorker(self.email, self.password,search_criteria,keyword,selected_date,selected_time,attachment_root_dir='my_attachments')
        self.worker.finished.connect(self.display_results)
        self.worker.start()

    def display_results(self, result):
        mail_ids, emails = result
        self.emails = emails
        self.mail_ids = mail_ids

        self.email_list.clear()
        for email in emails:
            self.email_list.addItem(email['subject'])

    def display_email(self, item):
        index = self.email_list.row(item)
        email = self.emails[index]
        display_text = f"主題: {email['subject']}\n"
        display_text += f"來自: {email['from']}\n"
        display_text += f"收件人: {email['to']}\n"
        display_text += f"日期: {email['date']}\n"
        display_text += f"內容: {email['body']}\n"
        if email['attachments']:
            display_text += f"附件: {', '.join(email['attachments'])}\n"
        display_text += "\n" + "-" * 40 + "\n\n"

        self.email_display.setPlainText(display_text)
    def order2word(self):
        order2word_emails=self.emails
        order2word_mail_ids=self.mail_ids
        if not order2word_emails:
           order2word_dlg=order2wordDialog()
           print("click")
           if order2word_dlg.exec():
                print("Success!")
           else:
               print("Cancel!")

        if order2word_emails:
            print("這邊要處理把信箱訂單要的資訊轉到試算表或是整理成word，主要就是要統整訂單資料")
            order2word_word=""""""
            order2word_word_t=""""""
            attachment_root_dir="order2word"
            if not os.path.exists(attachment_root_dir):
                os.makedirs(attachment_root_dir)

            for email in order2word_emails:
                order2word_1=f"{email['body']}"
                # 判斷 body 是否包含 "大榮"
                if "大榮" in order2word_1:
                    order2word_2=order2word_1[order2word_1.find('附件為'):]
                    #####下面是處理訂單訊息的
                    start1 = order2word_2.find("附件為") + len("附件為")
                    end1 = order2word_2.find("訂單")

                    # 找到 "明細如下:" 和 最後一個 "箱" 的位置，並擷取中間的部分
                    start2 = order2word_2.find("明細如下:")+len("明細如下:")
                    end2 = order2word_2.rfind("箱") + len("箱")

                    # 確保找到的索引是有效的
                    if start1 != -1 and end1 != -1 and start2 != -1 and end2 != -1:
                        new_str = order2word_2[start1:end1]+"大榮-久富餘-菁弘" + order2word_2[start2:end2]+"\n"
                        order2word_word+=new_str
                    else:
                        print("找不到指定的子字元串")
                elif "通盈" in order2word_1:
                    print("處理通盈的")
                    order2word_2=order2word_1[order2word_1.find('附件為'):]
                    #####下面是處理訂單訊息的
                    start1 = order2word_2.find("附件為") + len("附件為")
                    end1 = order2word_2.find("訂單")

                    # 找到 "明細如下:" 和 最後一個 "箱" 的位置，並擷取中間的部分
                    start2 = order2word_2.find("明細如下:")+len("明細如下:")
                    end2 = order2word_2.rfind("箱") + len("箱")

                    # 確保找到的索引是有效的
                    if start1 != -1 and end1 != -1 and start2 != -1 and end2 != -1:
                        new_str = order2word_2[start1:end1]+"通盈-久富餘-菁弘" + order2word_2[start2:end2]+"\n"
                        order2word_word_t+=new_str
                    else:
                        print("找不到指定的子字元串")

            if  order2word_word:
                print(f"要做分行跟關鍵字提取，這是全部訂單內容:\n{order2word_word}")   
                ####下面是把訂單資料存到word        
                current_date=self.order_search_datetime#獲取日期"1030602"是這個格式
                file_name = f"{current_date}大榮.docx"  # 生成文件名
                # 完整文件路径
                file_path = os.path.join(attachment_root_dir, file_name)
                # 创建一个新的 Word 文档
                document = Document()
                for line in order2word_word.split('\n'):
                    para = document.add_paragraph()  # 创建一个新段落
                    run = para.add_run(line)  # 在段落中添加一个运行并设置文本
                    run = para.runs[0]
                    run.font.size = Pt(16)  # 设置字体大小
                    run.font.name = 'Arial'  # 设置字体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')# 设置中文字体
                # 保存文档
                document.save(file_path)
                print(f"Word 文档已保存至: {file_path}\n")

            if order2word_word_t:
                print(f"要做分行跟關鍵字提取，這是全部訂單內容:\n{order2word_word_t}")   
                ####下面是把訂單資料存到word        
                current_date=self.order_search_datetime#獲取日期"1030602"是這個格式
                print(f"{current_date}")
                file_name = f"{current_date}通盈.docx"  # 生成文件名
                # 完整文件路径
                file_path_1 = os.path.join(attachment_root_dir, file_name)
                # 创建一个新的 Word 文档
                document = Document()
                for line in order2word_word_t.split('\n'):
                    para = document.add_paragraph()  # 创建一个新段落
                    run = para.add_run(line)  # 在段落中添加一个运行并设置文本
                    run = para.runs[0]
                    run.font.size = Pt(16)  # 设置字体大小
                    run.font.name = 'Arial'  # 设置字体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')# 设置中文字体
                # 保存文档
                document.save(file_path_1)
                print(f"Word 文档已保存至: {file_path_1}\n")

    def send_email(self):
        print("寄送郵件功能")
        send_email_emails=self.emails
        send_email_mail_ids=self.mail_ids
        if not send_email_emails:
           order2word_dlg=order2wordDialog()
           print("click")
           if order2word_dlg.exec():
                print("Success!")
           else:
               print("Cancel!")
        from_address = self.from_input_1.text()
        if not from_address:
            order2word_dlg=send_emailDialog()
            print("click")
            if order2word_dlg.exec():
                print("Success!")
            else:       
               print("Cancel!")
        

            


class EmailWorker(QThread):
    finished = pyqtSignal(tuple)

    def __init__(self, email, password,search_criteria,keyword,selected_date,selected_time,attachment_root_dir='attachments'):
        super().__init__()
        self.email = email
        self.password = password
        self.search_criteria = search_criteria
        self.keyword=keyword
        self.attachment_root_dir = attachment_root_dir
        self.selected_date=selected_date
        self.selected_time=selected_time
                # 检查并删除已存在的目录
        if os.path.exists(self.attachment_root_dir):
            shutil.rmtree(self.attachment_root_dir)
        if not os.path.exists(self.attachment_root_dir):
            os.makedirs(self.attachment_root_dir)

    def decode_mime_words(self, s):
        decoded_bytes, charset = decode_header(s)[0]
        if charset:
            decoded_str = decoded_bytes.decode(charset)
        else:
            decoded_str = decoded_bytes
        return decoded_str


    def run(self):
        try:
            mail = imaplib.IMAP4_SSL('imap.gmail.com')
            mail.login(self.email, self.password)
            mail.select('inbox')

            search_criteria = f'{self.search_criteria}'
            print(f"這邊是條件:{search_criteria}")
            result, data = mail.search(None, search_criteria)
            if result != 'OK':
                print("搜索邮件失败")
                self.finished.emit([])
                return

            mail_ids = data[0].split()
            print(f"{mail_ids}")
            print(f"{len(mail_ids)}")
            emails = []
            for i in range(len(mail_ids)): #range(len(mail_ids) - 1, -1, -1):
                result, data = mail.fetch(mail_ids[i], '(RFC822)')
                raw_email = data[0][1]
                msg = email.message_from_bytes(raw_email)
                # 获取邮件的日期
                date_tuple = email.utils.parsedate_tz(msg['Date'])
                if date_tuple:
                    local_date = datetime.fromtimestamp(email.utils.mktime_tz(date_tuple))
                    print(f"邮件日期: {local_date}")

                    # 进一步过滤到特定时间范围
                    selected_date=self.selected_date
                    start_time = datetime.strptime(selected_date + " 00:00", "%d-%b-%Y %H:%M")
                    end_time = datetime.strptime(selected_date + f" {self.selected_time}", "%d-%b-%Y %H:%M")

                    if start_time <= local_date <= end_time:
                        # 处理符合条件的邮件
                        print(f"处理符合时间范围的邮件: {local_date}")
                        email_data = {
                        'subject': self.decode_mime_words(msg['subject']),
                        'from': self.decode_mime_words(msg['from']),
                        'to': msg['to'],
                        'date': msg['date'],
                        'body': '',
                        'attachments': []
                        }
                        # 获取邮件的日期并格式化
                        email_date = msg.get('Date')
                        date_obj = email.utils.parsedate_to_datetime(email_date)
                        formatted_date = date_obj.strftime('%Y-%m-%d')+"_"+str(mail_ids[i])

                        # 创建以日期命名的文件夹
                        attachment_dir = os.path.join(self.attachment_root_dir, formatted_date)
                        if not os.path.exists(attachment_dir):
                            os.makedirs(attachment_dir)


                        if msg.is_multipart():
                            for part in msg.walk():
                                content_type = part.get_content_type()
                                content_disposition = str(part.get('Content-Disposition'))

                                if "attachment" in content_disposition:
                                    filename = part.get_filename()
                                    if filename:
                                        decoded_filename = self.decode_mime_words(filename)
                                        decoded_filename = decoded_filename.replace("/", "_")  # 确保文件名合法
                                        # 将附件保存到指定目录
                                        attachment_path = os.path.join(attachment_dir, decoded_filename)
                                        email_data['attachments'].append(attachment_path)
                                        with open(attachment_path, 'wb') as f:
                                            f.write(part.get_payload(decode=True))
                                elif content_type == "text/plain":
                                    email_data['body'] = part.get_payload(decode=True).decode('utf-8')
                        else:
                            email_data['body'] = msg.get_payload(decode=True).decode('utf-8')

                        emails.append(email_data)

            mail.logout()

            keyword = self.keyword
            filtered_emails = [email for email in emails if keyword in email['body']]
            # 獲取過濾後的郵件的索引
            filtered_indices = [index for index, email in enumerate(emails) if keyword in email['body']]

            # 獲取對應的 mail_ids
            filtered_mail_ids = [mail_ids[index] for index in filtered_indices]

            # 獲取過濾後的郵件
            filtered_emails = [emails[index] for index in filtered_indices]
            print(f"{filtered_mail_ids}")
            self.finished.emit((filtered_mail_ids, filtered_emails))
            

        except Exception as e:
            print(f"Error: {e}")
            self.finished.emit(())



class order2wordDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("阿肥之力")
        self.setMinimumSize(180,80)
        self.setMaximumSize(180,80)

        QBtn = QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel

        self.buttonBox = QDialogButtonBox(QBtn)
        self.buttonBox.accepted.connect(self.accept)
        self.buttonBox.rejected.connect(self.reject)

        self.layout = QVBoxLayout()
        message = QLabel("沒有搜索資料，請先搜索")
        self.layout .addWidget(message)
        self.layout .addWidget(self.buttonBox)
        self.setLayout(self.layout)

class send_emailDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("阿肥之力")
        self.setMinimumSize(180,80)
        self.setMaximumSize(180,80)

        QBtn = QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel

        self.buttonBox = QDialogButtonBox(QBtn)
        self.buttonBox.accepted.connect(self.accept)
        self.buttonBox.rejected.connect(self.reject)

        self.layout = QVBoxLayout()
        message = QLabel("請輸入收件人郵件")
        self.layout .addWidget(message)
        self.layout .addWidget(self.buttonBox)
        self.setLayout(self.layout)
